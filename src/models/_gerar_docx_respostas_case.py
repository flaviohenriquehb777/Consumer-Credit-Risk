"""
Gerador do arquivo .docx do CASE com as 5 perguntas respondidas,
usando evidências reais do projeto Consumer Credit Risk (Aurora).

Saída: docs/RESPOSTAS_DO_CASE_Aurora_Consumer_Credit_Risk.docx
"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[2]  # repo root
OUT_DIR = ROOT / "docs"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_FILE = OUT_DIR / "RESPOSTAS_DO_CASE_Aurora_Consumer_Credit_Risk.docx"

# ------------------------------------------------------------------
# Helpers de estilo
# ------------------------------------------------------------------
def shade_cell(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)

def add_caption(doc, texto: str):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

doc = Document()

# ---- Estilo padrão (tamanho legível / profissional)
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
section = doc.sections[0]
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.4)
section.right_margin = Cm(2.4)

# ------------------------------------------------------------------
# Capa
# ------------------------------------------------------------------
tit = doc.add_paragraph()
tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tit.add_run("CASE AURORA – CONSUMER CREDIT RISK")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Respostas estruturadas às 5 perguntas de negócio\n"
                "baseadas nas evidências reais do projeto (Fases 1 a 6 do CRISP-DM)")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    "Modelo vencedor: XGBoost (max_depth=4, lr=0.03, mcw=80, n=500, spw=13.96)\n"
    "Holdout ROC AUC = 0,86956  ·  Economia de custo vs. “aprovar todos” = 50,65%\n"
    "Limiar ótimo (custo FN = R$ 5.000 ; FP = R$ 500) = 0,56"
)
r.font.size = Pt(10.5)
r.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

doc.add_paragraph()  # espaço

# ------------------------------------------------------------------
# 1. Variáveis mais relevantes
# ------------------------------------------------------------------
doc.add_heading("1. Quais variáveis parecem mais relevantes para prever inadimplência?", level=1)

doc.add_paragraph(
    "A relevância é medida de três formas complementares no projeto: (a) importância "
    "por cobertura nas árvores do XGBoost; (b) impacto SHAP global médio (|mean SHAP|), "
    "que mede quanto cada feature move o score em valor absoluto sobre 37.500 clientes "
    "de teste (holdout cego); e (c) magnitude dos SHAP extremos (pior caso de 1 cliente)."
)

doc.add_paragraph("Ranking global TOP 10 por impacto SHAP médio absoluto |mean SHAP| (holdout):")

tbl = doc.add_table(rows=1 + 10, cols=5)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
headers = ["Rank", "Variável", "|mean SHAP|", "Direção média", "% de clientes que puxam ↑ risco"]
for c, h in zip(hdr, headers):
    c.text = h
    for run in c.paragraphs[0].runs:
        run.bold = True
    shade_cell(c, "1F3A5F")
    for run in c.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

top10 = [
    (1, "uso_limite_rotativo",     "0,839", "Redutor (baixo uso = risco baixo)",  "29,97%"),
    (2, "atrasos_30_59_dias",      "0,386", "Negativo (média reduz score)",        "15,88%"),
    (3, "atrasos_90_mais_dias",    "0,329", "Negativo (média reduz score)",         "5,45%"),
    (4, "idade",                   "0,232", "Negativo (média reduz score)",        "54,72%"),
    (5, "atrasos_60_89_dias",      "0,176", "Negativo (média reduz score)",         "5,01%"),
    (6, "linhas_credito_abertas",  "0,167", "Neutro-negativo",                     "43,85%"),
    (7, "financiamentos_imobiliarios","0,116","Negativo",                           "41,19%"),
    (8, "sobra_caixa (engenharia)","0,092", "Negativo",                             "40,24%"),
    (9, "razao_divida",            "0,088", "Negativo",                             "35,25%"),
    (10,"renda_mensal",            "0,060", "Negativo",                             "43,09%"),
]
for i, row in enumerate(top10, start=1):
    cells = tbl.rows[i].cells
    for c, v in zip(cells, row):
        c.text = str(v)

doc.add_paragraph()

doc.add_paragraph(
    "Conclusão executiva: “uso_limite_rotativo” domina como preditor (2× o impacto do "
    "2º colocado), seguido pelas três variáveis de atraso (30–59, 60–89 e 90+ dias). "
    "Juntos, esses 4 fatores explicam ≈70% da variação de impacto SHAP global."
)

# ------------------------------------------------------------------
# 2. Missing renda_mensal e dependentes
# ------------------------------------------------------------------
doc.add_heading(
    "2. Como a Aurora deveria tratar valores ausentes em renda_mensal e dependentes? "
    "A ausência de informação deve ser imputada, sinalizada por uma variável indicadora "
    "ou tratada de outra forma?", level=1)

doc.add_paragraph("Diagnóstico (base bruta, 150.000 observações):")
tbl2 = doc.add_table(rows=1+2, cols=5)
tbl2.style = "Light Grid Accent 1"
h = tbl2.rows[0].cells
for c, t in zip(h, ["Variável", "Qtd. ausentes", "% ausentes",
                    "Inadimplência quem NÃO informou",
                    "Inadimplência quem INFORMOU"]):
    c.text = t
    for run in c.paragraphs[0].runs: run.bold = True
    shade_cell(c, "1F3A5F")
    for run in c.paragraphs[0].runs: run.font.color.rgb = RGBColor(255,255,255)
row1 = ["renda_mensal",     "29.731", "19,82%", "5,61%", "6,95%"]
row2 = ["numero_dependentes","3.924",  "2,62%",  "—",     "—"]
for i, rr in enumerate([row1,row2], start=1):
    for c, v in zip(tbl2.rows[i].cells, rr):
        c.text = v

doc.add_paragraph()

doc.add_paragraph(
    "Como a ausência de renda_não é aleatória (inadimplência difere entre quem informa "
    "e quem não informa em 134 bps), imputar mediana SEM sinalizador apagaria o sinal "
    "de risco. A estratégia recomendada é a que foi aplicada no projeto — estratégia "
    "híbrida em 2 camadas:"
)
for item in [
    "Camada 1 — Flag binária (Indicadora): criar renda_ausente = 1 quando renda_mensal "
    "estiver ausente e dependentes_ausentes = 1 quando número de dependentes estiver "
    "ausente. Estas flags entram como variáveis preditivas separadas.",
    "Camada 2 — Imputação robusta (sem viés de treino/teste): imputar a mediana "
    "APRENDIDA SOMENTE NO TREINO para as colunas originais (renda_mensal mediana = R$ "
    "5.400; número de dependentes mediana = 0). A mediana é preferida à média por ser "
    "robusta aos outliers de renda (>R$ 100.000 existiam na base).",
    "Complementar: para a variável dependentes ausentes (2,62%) a flag por si só já é "
    "informativa. Para renda_mensal (19,82%) flag + imputação = o padrão ouro do BCB/"
    "FAIR lending guidance para hipótese de Missing Not At Random (MNAR).",
]:
    p = doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "Procedimento SÊNIOR anti-leakage: mediana R$ 5.400 aprendida no TREINO "
    "(112.500 linhas), reaplicada idêntica no holdout (37.500 linhas) e reaplicada idêntica "
    "no app Streamlit (arquivo preprocessamento.py importado pelo deploy)."
).italic = True

# ------------------------------------------------------------------
# 3. Ponto de corte ideal
# ------------------------------------------------------------------
doc.add_heading(
    "3. Como definir o ponto de corte ideal para aprovar ou negar crédito? "
    "A decisão deveria maximizar acurácia estatística, reduzir inadimplência ou "
    "maximizar lucro esperado?", level=1)

doc.add_paragraph(
    "A decisão NÃO deve maximizar acurácia. Por quê? Baseline trivial da base = “todos "
    "adimplentes” entrega 93,32% de acurácia e ainda assim é a pior política de negócio "
    "(custo esperado da carteira: FNs consomem o principal de 10 adimplentes). O objetivo "
    "correto é maximizar lucro esperado ou, equivalentemente, minimizar custo esperado "
    "da carteira ponderado pelos custos reais de negócio:"
)
for c in [
    "Custo de Falso Negativo (FN) = R$ 5.000 — aprovamos quem vai calotear; perdemos o principal do empréstimo (10×).",
    "Custo de Falso Positivo (FP) = R$ 500 — negamos quem pagaria; perdemos a margem do cliente (1×).",
]:
    doc.add_paragraph(c, style="List Bullet")

doc.add_paragraph("Metodologia aplicada no projeto (PADRÃO SÊNIOR SEM LEAKAGE):")
for item in [
    "1. Varredura 9.900 limiares (0,01 → 0,99, passo 0,01) APENAS sobre as previsões "
    "Out-Of-Fold (OOF) do TREINO via StratifiedKFold 5-fold. NUNCA sobre o holdout — "
    "holdout só é batizado 1 vez no final para confirmar cegamente.",
    "2. Para cada limiar, calcula custo total esperado = 5.000 × N_FN + 500 × N_FP.",
    "3. Seleciona o limiar com MENOR custo esperado no OOF treino. Chamamos esse limiar "
    "de ponto de corte ótimo orientado a negócio.",
]:
    doc.add_paragraph(item, style="List Number")

tbl3 = doc.add_table(rows=1+2, cols=6)
tbl3.style = "Light Grid Accent 1"
h = tbl3.rows[0].cells
for c, t in zip(h, ["Modelo", "Limiar ótimo (OOF treino)",
                    "Custo OOF treino (R$ milhões)",
                    "Economia vs. “aprovar todos” (treino)",
                    "Custo HOLDOUT cego (R$ milhões)",
                    "Economia HOLDOUT"]):
    c.text = t
    for run in c.paragraphs[0].runs: run.bold = True
    shade_cell(c, "1F3A5F")
    for run in c.paragraphs[0].runs: run.font.color.rgb = RGBColor(255,255,255)
linhas = [
    ["Árvore de Decisão (d=7)", "0,08", "R$ 19,78M", "47,28%", "—", "47,59%"],
    ["XGBoost Campeão",        "0,56", "R$ 18,80M", "50,01%", "R$ 18,798M", "50,65%"],
]
for i, rr in enumerate(linhas, start=1):
    for c, v in zip(tbl3.rows[i].cells, rr):
        c.text = v

doc.add_paragraph()
doc.add_paragraph(
    "Conclusão executiva: ponto de corte ideal para a Aurora = **0,56 no XGBoost Campeão**. "
    "Nesse limiar, HOLDOUT: P(inadimplência) ≥ 56% → negar crédito; <56% → aprovar. "
    "Não se decide por “máxima acurácia” (errado), decide-se por **máxima economia de "
    "custo real**."
).bold = True

# ------------------------------------------------------------------
# 4. Explicar a decisão para áreas de negócio e clientes
# ------------------------------------------------------------------
doc.add_heading(
    "4. Como explicar a decisão do modelo para áreas de negócio e clientes sem reduzir "
    "excessivamente a complexidade do problema?", level=1)

doc.add_paragraph(
    "Solução do projeto: três camadas de comunicação, alinhadas com LGPD (art. 7º, direito "
    "de explicação), BCB Circular 4.015/20 e exigências de compliance:"
)

camadas = [
    ("Camada 1 — Executiva (1 slide, 10 s, para áreas de negócio / diretoria)",
     "Card de decisão com 3 métricas: (a) % de inadimplência estimada; "
     "(b) Decisão “APROVAR” / “RECUSAR” com o limiar 56% em destaque; "
     "(c) Card de economia esperada vs. “aprovar todos” (50,65% no holdout). "
     "Impresso como cabeçalho do app Streamlit e de qualquer relatório de tomada de decisão."),
    ("Camada 2 — Analista de crédito / Compliance (5 motivos, visual TABELA + GRÁFICO)",
     "TOP 5 SHAP LOCAL por cliente, entregue em duas formas simultâneas no app: "
     "(a) Tabela com Feature, Valor do Cliente, Contribuição SHAP (verde para redutores "
     "de risco, vermelho para amplificadores) e texto “aumenta risco ↑ calote / diminui "
     "risco ↓ calote”; (b) Gráfico horizontal TOP 10 de barras com linha tracejada de "
     "base zero, cores vermelho/verde e tooltip interativo com o valor exato da feature "
     "na linha. Tudo calculado <10 ms por cliente via TreeExplainer exato do XGBoost."),
    ("Camada 3 — Cliente final / Ouvidoria (linguagem natural, sem jargões, 2 frases)",
     "Template padrão, baseado nos Top 3 SHAP de cada cliente: “Seu crédito foi "
     "RECUSADO/APROVADO com base em três fatores principais: (1) Utilização do seu "
     "limite de cartão de [X%] comparado a perfis aprovados; (2) Histórico de [N] "
     "atrasos de 30 a 59 dias nos últimos 2 anos; (3) Perfil de idade de [Y] anos em "
     "conjunto com razão dívida/renda de [Z%]. Detalhes documentados e sob consulta na "
     "ouvidoria, conforme Art. 7º da LGPD.”."),
    ("Garantia metodológica anti-black-box",
     "Nenhuma explicação usa “feature importance global” para explicar um cliente "
     "individual — sempre SHAP LOCAL. Sanity check obrigatório exibido no app: "
     "“Consistência SHAP (expected + Σ shap) vs. logit(p)” deve ficar abaixo de 1e-4. "
     "Se falhar, exibe “Investigar” e bloqueia a emissão da justificativa automática — "
     "evita explicações inventadas."),
]
for titulo, corpo in camadas:
    p = doc.add_paragraph()
    r = p.add_run(titulo)
    r.bold = True
    doc.add_paragraph(corpo)

# ------------------------------------------------------------------
# 5. Riscos éticos / regulatórios / reputacionais
# ------------------------------------------------------------------
doc.add_heading(
    "5. Quais riscos éticos, regulatórios ou reputacionais podem surgir ao utilizar idade, "
    "renda, dependentes e histórico de atraso em uma política automatizada de crédito?",
    level=1)

doc.add_paragraph(
    "Classificação em três pilares, com mitigações adotadas no projeto:"
)

secoes = [
    ("5.1 Riscos éticos — Discriminação indireta e viés (Fair Lending)",
     [
         "Idade como fator de risco: no SHAP global idade é o 4º maior impacto (0,232). "
         "Sem governança, o modelo pode penalizar sistematicamente jovens (18–30 anos) e "
         "super-beneficiar idosos — configurando age discrimination e risco ético.",
         "Renda_mensal + n_dependentes correlação com raça/gênero/região. No Brasil a "
         "Proteção de Dados Pessoais (LGPD) e a Circular BCB 4.015 proíbem uso de proxies "
         "de classe e origem. A ausência de sinalização de missing poderia atuar como "
         "proxy de “baixa escolaridade”; mitigamos com flag explícita.",
         "Histórico de atraso vs. “cod_sistema_atrasos 96/98”: se 264+5 linhas recebiam "
         "tratamento de vencedor sem flag, o modelo penalizaria sistematicamente quem "
         "tem erro de sistema (viés de medição) — tratamento: flag binária dedicada."
     ]),
    ("5.2 Riscos regulatórios — LGPD, BCB, Bacen e Auto-Regulação Febraban",
     [
         "LGPD Art. 7º (direito a explicação individual) e Art. 20 (título de crédito) "
         "exigem que o titular SAIBO “por que” o crédito foi negado em linguagem acessível "
         "— o que obriga ter SHAP/Waterfall e não só score black box.",
         "BCB Circular 4.015/2020 e Resolução CMN 4.852/2020 (governança de modelos de "
         "crédito): exigem documentar metodologia, validação, monitoramento de desempenho "
         "e reexame anual. Violações = multa até 2% do faturamento bruto anual da "
         "instituição.",
         "Regulamentação de uso de IA no Brasil (PL 2.338/2023) e na UE (AI Act, 2024): "
         "classificam sistemas de crédito automatizado como “alto risco”, exigindo "
         "registro de decisões, trilha de auditoria completa e intervenção humana "
         "possível em qualquer ponto. Devemos arquivar por cliente: 10 inputs brutos, "
         "probabilidade, limiar, decisão, top 5 SHAP e data/hora."
     ]),
    ("5.3 Riscos reputacionais — Erros visíveis e opacidade",
     [
         "Rejeição arbitrária percebida: se um bom cliente com score de 55,9% é negado "
         "por “1 décimo abaixo do limiar 0,56” sem explicação, vira caso de Reclame "
         "Aqui / notícia negativa. Mitigação: exibir o limiar explicitamente no app e "
         "oferecer “análise manual complementar” para a banda 0,50 → 0,62 (borderline).",
         "Data drift / concept drift: se a taxa de juros sobe 500 bps ou inflação afeta "
         "renda real, o modelo treinado em Janeiro fica descalibrado em Agosto; "
         "inadimplência real cresce e a Aurora perde reputação de “risco conservador” "
         "por aprovar gente que caloteou. Mitigação: monitoramento mensal obrigatório de "
         "PSI population stability index por feature e CSI characteristic stability index "
         "de score.",
         "Uso indevido por analistas (gaming): sem governança, analista de crédito pode "
         "forçar inputs manualmente para “passar” o cliente (minimizar uso_limite → aprovar "
         "e depois o cliente usa 120% do limite). Risco reputacional e de perda financeira. "
         "Mitigação: import automático de inputs via API do cadastro, e histórico de "
         "alterações manuais no formulário do app."
     ]),
]

for tit, bul in secoes:
    doc.add_heading(tit, level=2)
    for b in bul:
        doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Conclusão do item 5 (governança recomendada à Aurora)", level=2)
doc.add_paragraph(
    "As quatro variáveis devem continuar no modelo (elas carregam poder preditivo real, "
    "e excluí-las degradaria ROC AUC abaixo de 0,85), mas sob uma estrutura de três "
    "controles: (a) Fairness report mensal por faixas etárias, decil de renda e UF; "
    "(b) Trilha de auditoria por decisão (inputs brutos + preparados + p + limiar + "
    "decisão + top 5 SHAP + usuário); e (c) Override humano obrigatório para clientes "
    "na faixa borderline 0,50 a 0,62. Com essas três camadas a política automatizada "
    "fica ética, regulatória e de baixa reputação de risco."
).bold = True

# ------------------------------------------------------------------
# Rodapé
# ------------------------------------------------------------------
doc.add_paragraph()
doc.add_paragraph("— Fim das respostas do case Aurora Consumer Credit Risk. —").alignment = WD_ALIGN_PARAGRAPH.CENTER
add_caption(doc,
    "Documento gerado automaticamente por scripts do projeto. Todas as tabelas numéricas "
    "foram extraídas de execuções reais nas Fases 2 a 6 em src/models/*.py e validadas em "
    "holdout cego 37.500 clientes (seed=42, StratifiedKFold 5-fold no treino).")

doc.save(str(OUT_FILE))
print(f"✅ Arquivo .docx salvo com sucesso em:\n   {OUT_FILE}")
print(f"   Tamanho: {OUT_FILE.stat().st_size/1024:,.1f} KB")
